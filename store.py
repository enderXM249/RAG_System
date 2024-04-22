import os
import sys
import streamlit as st
from dotenv import load_dotenv
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone, ServerlessSpec
import pdfplumber
import docx
from pydantic import ValidationError

# Function to extract text from PDF
def extract_text_from_pdf(uploaded_file):
    with pdfplumber.open(uploaded_file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise ValidationError("Error extracting text from DOCX file: " + str(e))


def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=20)
    chunks = text_splitter.split_text(text)
    return chunks

embedding = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")

def get_vector_store(docs):
    pc = Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))
    index_name = 'pinecone'
    # Convert the text chunks to a list of Document objects
    documents = [Document(page_content=chunk) for chunk in docs]
    PineconeVectorStore.from_documents(documents, embedding, index_name=index_name)
   